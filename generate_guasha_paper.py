#!/usr/bin/env python3
"""Generate Word paper: 刮痧 + 文明交流互鉴. Format: title SimHei 二号, body SimSun 四号, 1.5 line spacing."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


TITLE = "文化交流互鉴视域下的电影《刮痧》及时代启示"

# 正文（含标点）不少于1500字
BODY_PARAS = [
    "摘要：电影《刮痧》以一起因传统疗法引发的跨文化误解为叙事核心，集中呈现了不同文明传统在相遇时的认知落差与制度张力。本文在梳理影片情节与主题的基础上，结合习近平总书记关于文明交流互鉴的重要论述，从平等对话、相互理解与开放包容等维度，阐释文化交流互鉴对于个人交往、社会治理与人类文明进步的重要意义，并就新时代讲好中国故事、推动实践层面的文明对话提出思考。",
    "关键词：刮痧；文明交流互鉴；文化差异；平等对话",
    "一、引言",
    "全球化深入发展的今天，人员、信息与价值的跨国流动日益频繁，文化在交流互鉴中发展既是历史规律，也是现实课题。习近平总书记指出：“文明因多样而交流，因交流而互鉴，因互鉴而发展。”这一论断揭示了文明演进的内在逻辑：多样性是前提，交流是路径，互鉴是方法，发展是目标。电影《刮痧》虽创作于二十余年前，却以戏剧化的方式把“同一事实、不同解释”的跨文化困境推到观众面前，为理解“为何需要交流、何以实现互鉴”提供了富有张力的艺术案例。",
    "二、电影《刮痧》的主要内容与观后思考",
    "《刮痧》由郑晓龙执导，讲述旅美华人许大同一家在美国生活、奋斗的故事。许父从国内来美探亲，见孙子丹尼斯身体不适，便以中国传统医学中的刮痧为其调理。刮痧后皮肤出现的瘀痕，在缺乏文化语境解释的情况下，被美国社会按照“儿童受虐”的法律与医学框架加以解读，家庭一度陷入监护权纠纷与信任危机。影片通过法庭辩论、邻里关系与代际观念的交织，把中医经验传承、家庭伦理与西方法治话语并置在同一叙事空间，使观众直观感受到：当一种生活方式与知识体系“缺席于对方的词典”时，善意也可能被误读，传统也可能被误伤。",
    "从艺术表达看，《刮痧》的价值不在于简单褒贬某一种文化，而在于揭示差异的客观性与对话的必要性。影片中，许大同强调“亲情与孝”，美国社会强调“程序与证据”；中医讲究整体观与经验积累，西医与司法体系更依赖可重复验证与可见创伤的对应关系。冲突的根源并非“谁更先进”，而是缺少把不同知识体系转译为彼此可理解语言的中间环节。观后令人深思的是：若能在事前建立更充分的健康沟通、社区教育与多元文化服务，许多误会本可在制度与日常层面被提前化解；而当误会已经发生，真正能够弥合裂痕的，往往不是单方面的辩解，而是基于尊重的事实澄清、耐心倾听与共同寻求儿童最佳利益的合作。",
    "三、学习习近平总书记关于文化在交流互鉴中发展的重要论述",
    "习近平总书记立足人类前途命运与中华文明发展大势，系统阐发了平等、互鉴、对话、包容的文明观。总书记强调，人类文明因多样才有交流互鉴的价值，因平等才有交流互鉴的前提，因包容才有交流互鉴的动力，因对话才有交流互鉴的常态。总书记还以生动比喻指出，文明如果长期自我封闭，必将走向衰落；中华文明在长期历史进程中与其他文明不断交流互鉴，形成了开放包容的品格。这些重要论述为我们把握文化发展规律提供了根本遵循：尊重差异不是放任误解，而是在差异中寻找共识；推动互鉴不是消解自我，而是在对话中实现各美其美、美美与共。",
    "把论述落到实践层面，至少包含三层含义。其一，坚持平等对话，反对文明优越论与单一标准论，承认不同民族在长期实践中形成的知识传统与生活方式都有其历史合理性。其二，坚持开放包容，在守正创新中推动中华优秀传统文化创造性转化、创新性发展，以自信从容的姿态参与世界文明百花园的建设。其三，坚持问题导向，针对跨文化交往中的信息不对称与结构性偏见，完善传播方式、公共叙事与制度性安排，使交流从“偶发解释”走向“常态机制”。",
    "四、结合实际谈对文化交流互鉴重要性的认识",
    "结合当下实际，文化交流互鉴的重要性至少体现在三个方面。",
    "第一，在个体与家庭层面，互鉴关乎尊严与幸福。无论是留学、务工、商务还是移民家庭，跨文化处境中的人们常常像《刮痧》的主人公一样，需要在“解释自己”与“理解他者”之间不断切换。只有以平等心态学习对方社会的规则与表达，同时善于用准确、温和、可验证的方式介绍自身传统，才能减少不必要的摩擦，让不同背景的人真正“看见彼此”。",
    "第二，在社会与治理层面，互鉴关乎信任与效能。公共服务、医疗教育、社区警务、媒体传播等领域，若能引入跨文化培训、多语种信息与多元参与机制，就能降低因认知落差导致的误判成本。《刮痧》给我们的提醒是：法律与程序正义极为重要，但正义的实现同样依赖文化理解与沟通艺术；把“差异”纳入治理视野，并不是削弱原则，而是让原则更具可及性与可接受性。",
    "第三，在文明与发展层面，互鉴关乎创新与人类共同利益。应对气候变化、公共卫生、人工智能治理等全球性议题，需要不同文明贡献智慧资源。文化交流越充分，越能在相互启发中形成解决问题的新方案；文明对话越深入，越能在共同挑战面前凝聚合作意愿。中国提出全球文明倡议等重大主张，正是对“交流互鉴促发展”的时代回应。",
    "五、结语",
    "电影《刮痧》以一场误会照见跨文化交往的复杂性与脆弱性，也反向证明了交流互鉴的不可替代性。学习习近平总书记关于文明交流互鉴的重要论述，我们更应坚定文化自信，同时以海纳百川的胸怀推动对话，在尊重法治与人类共同价值的前提下，讲好中国故事、传播好中国声音，让世界在真实、立体、全面的认知中理解中国。惟其如此，文明才能在互鉴中生生不息，人类才能在差异中携手走向更美好的未来。",
]


def set_run_font(run, name_cn: str, name_west: str, size_pt: float):
    run.font.name = name_west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_cn)
    run.font.size = Pt(size_pt)


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)

    # 标题：黑体 二号（22磅）
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(18)
    p0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p0.paragraph_format.line_spacing = 1.5
    r0 = p0.add_run(TITLE)
    set_run_font(r0, "黑体", "SimHei", 22)
    r0.bold = False

    for text in BODY_PARAS:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(28)  # 约等于四号字2字符缩进
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        # 小标题行不缩进
        if text.startswith(("一、", "二、", "三、", "四、", "五、")) and len(text) < 30:
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(text)
            set_run_font(run, "宋体", "SimSun", 14)
            run.bold = True
            continue
        if text.startswith("关键词："):
            p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        set_run_font(run, "宋体", "SimSun", 14)

    full_text = TITLE + "\n" + "\n".join(BODY_PARAS)
    n = len(full_text.replace("\n", ""))
    out = "/Users/laip/Desktop/code/Laip11.github.io/刮痧与文化交流互鉴论文.docx"
    doc.save(out)
    print("saved:", out)
    print("char_count (no newlines):", n)


if __name__ == "__main__":
    main()
